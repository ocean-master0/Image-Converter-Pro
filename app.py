from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from PIL import Image
import os
import io
import base64
import uuid
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches
import svgwrite
import tempfile
import shutil
import time
import threading
from datetime import datetime

app = Flask(__name__)
app.secret_key = 'your-secret-key-here-change-in-production'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Create directories with proper error handling
def create_directories():
    try:
        upload_path = os.path.abspath(app.config['UPLOAD_FOLDER'])
        output_path = os.path.abspath(app.config['OUTPUT_FOLDER'])
        
        os.makedirs(upload_path, exist_ok=True)
        os.makedirs(output_path, exist_ok=True)
        
        app.config['UPLOAD_FOLDER'] = upload_path
        app.config['OUTPUT_FOLDER'] = output_path
        
        print(f"Directories created:")
        print(f"  Upload: {upload_path}")
        print(f"  Output: {output_path}")
        
    except Exception as e:
        print(f"Error creating directories: {e}")
        raise

create_directories()

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_old_files():
    try:
        current_time = time.time()
        deleted_count = 0
        
        for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
            if os.path.exists(folder):
                for filename in os.listdir(folder):
                    file_path = os.path.join(folder, filename)
                    if os.path.isfile(file_path):
                        try:
                            if current_time - os.path.getctime(file_path) > 3600:
                                os.remove(file_path)
                                deleted_count += 1
                        except Exception as e:
                            print(f"Error removing file {file_path}: {e}")
        
        if deleted_count > 0:
            print(f"Cleaned up {deleted_count} old files")
            
    except Exception as e:
        print(f"Error in cleanup: {e}")

def schedule_cleanup():
    cleanup_old_files()
    timer = threading.Timer(1800, schedule_cleanup)
    timer.daemon = True
    timer.start()

schedule_cleanup()

class ImageConverter:
    def __init__(self, image_path):
        try:
            self.image = Image.open(image_path)
            self.image_path = image_path
        except Exception as e:
            raise Exception(f"Failed to open image: {str(e)}")

    def convert_to_image(self, save_path, format_type):
        try:
            if format_type == "JPEG" and self.image.mode in ("RGBA", "P"):
                background = Image.new("RGB", self.image.size, (255, 255, 255))
                if self.image.mode == "RGBA":
                    background.paste(self.image, mask=self.image.split()[-1])
                else:
                    background.paste(self.image)
                background.save(save_path, "JPEG", quality=95, optimize=True)
            else:
                self.image.save(save_path, format_type, optimize=True)
        except Exception as e:
            raise Exception(f"Image conversion failed: {str(e)}")

    def convert_to_svg(self, save_path):
        try:
            img_buffer = io.BytesIO()
            if self.image.mode == "RGBA":
                self.image.save(img_buffer, format='PNG')
            else:
                rgb_image = self.image.convert('RGB')
                rgb_image.save(img_buffer, format='PNG')
            
            img_buffer.seek(0)
            img_base64 = base64.b64encode(img_buffer.getvalue()).decode()

            dwg = svgwrite.Drawing(save_path, size=self.image.size)
            dwg.add(dwg.image(
                href=f"data:image/png;base64,{img_base64}",
                insert=(0, 0),
                size=self.image.size
            ))
            dwg.save()
        except Exception as e:
            raise Exception(f"SVG conversion failed: {str(e)}")

    def convert_to_pdf(self, save_path):
        """Fixed PDF conversion method"""
        try:
            # Convert image to RGB if needed
            if self.image.mode != 'RGB':
                rgb_image = self.image.convert('RGB')
            else:
                rgb_image = self.image

            # Create a temporary file for the image
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_image_path = temp_file.name
                rgb_image.save(temp_image_path, 'PNG', optimize=True)

            try:
                # Create PDF
                c = canvas.Canvas(save_path, pagesize=letter)
                page_width, page_height = letter
                img_width, img_height = rgb_image.size

                # Calculate scaling to fit page with margins
                scale = min(page_width / img_width, page_height / img_height) * 0.8
                scaled_width = img_width * scale
                scaled_height = img_height * scale

                # Center the image on the page
                x = (page_width - scaled_width) / 2
                y = (page_height - scaled_height) / 2

                # Draw image using the temporary file path
                c.drawImage(temp_image_path, x, y, width=scaled_width, height=scaled_height)
                c.save()

            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_image_path)
                except:
                    pass

        except Exception as e:
            raise Exception(f"PDF conversion failed: {str(e)}")

    def convert_to_docx(self, save_path):
        try:
            # Convert to RGB if needed
            if self.image.mode != 'RGB':
                rgb_image = self.image.convert('RGB')
            else:
                rgb_image = self.image

            img_buffer = io.BytesIO()
            rgb_image.save(img_buffer, format='PNG', optimize=True)
            img_buffer.seek(0)

            doc = Document()
            doc.add_heading('Converted Image', 0)

            img_width, img_height = rgb_image.size
            max_width = 6.0

            if img_width > img_height:
                width = Inches(max_width)
                height = Inches((img_height / img_width) * max_width)
            else:
                height = Inches(max_width)
                width = Inches((img_width / img_height) * max_width)

            paragraph = doc.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(img_buffer, width=width, height=height)
            paragraph.alignment = 1

            doc.save(save_path)
        except Exception as e:
            raise Exception(f"DOCX conversion failed: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file selected'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if file and allowed_file(file.filename):
            filename = str(uuid.uuid4()) + '.' + file.filename.rsplit('.', 1)[1].lower()
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            file.save(filepath)

            try:
                with Image.open(filepath) as img:
                    image_info = {
                        'filename': filename,
                        'original_name': file.filename,
                        'size': img.size,
                        'format': img.format,
                        'mode': img.mode,
                        'file_size': os.path.getsize(filepath)
                    }
                return jsonify(image_info)
            except Exception as e:
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                return jsonify({'error': f'Invalid image file: {str(e)}'}), 400

        return jsonify({'error': 'Invalid file type. Supported: JPG, PNG, GIF, BMP, TIFF, WebP'}), 400

    except Exception as e:
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/convert', methods=['POST'])
def convert_image():
    try:
        data = request.get_json()
        filename = data.get('filename')
        format_type = data.get('format')

        if not filename or not format_type:
            return jsonify({'error': 'Missing filename or format'}), 400

        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(input_path):
            return jsonify({'error': 'File not found'}), 404

        try:
            converter = ImageConverter(input_path)
            base_name = filename.rsplit('.', 1)[0]
            output_filename = f"{base_name}.{format_type.lower()}"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

            format_upper = format_type.upper()
            if format_upper in ['JPG', 'JPEG']:
                converter.convert_to_image(output_path, 'JPEG')
            elif format_upper == 'PNG':
                converter.convert_to_image(output_path, 'PNG')
            elif format_upper == 'SVG':
                converter.convert_to_svg(output_path)
            elif format_upper == 'PDF':
                converter.convert_to_pdf(output_path)
            elif format_upper == 'DOCX':
                converter.convert_to_docx(output_path)
            else:
                return jsonify({'error': 'Unsupported format'}), 400

            if not os.path.exists(output_path):
                return jsonify({'error': 'Conversion failed - output file not created'}), 500

            return jsonify({
                'success': True,
                'download_url': f'/download/{output_filename}',
                'filename': output_filename,
                'file_size': os.path.getsize(output_path)
            })

        except Exception as e:
            return jsonify({'error': f'Conversion failed: {str(e)}'}), 500

    except Exception as e:
        return jsonify({'error': f'Request processing failed: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        filename = secure_filename(filename)
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        
        if os.path.exists(file_path):
            def delete_after_delay():
                time.sleep(5)
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        print(f"Deleted downloaded file: {filename}")
                except Exception as e:
                    print(f"Error deleting file {filename}: {e}")
            
            threading.Thread(target=delete_after_delay, daemon=True).start()
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            return jsonify({'error': 'File not found'}), 404
            
    except Exception as e:
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

@app.route('/clear-all', methods=['POST'])
def clear_all_data():
    try:
        deleted_files = 0
        errors = []
        
        for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
            if os.path.exists(folder):
                for filename in os.listdir(folder):
                    file_path = os.path.join(folder, filename)
                    if os.path.isfile(file_path):
                        try:
                            os.remove(file_path)
                            deleted_files += 1
                        except Exception as e:
                            errors.append(f"Failed to delete {filename}: {str(e)}")
        
        response = {
            'success': True,
            'message': f'Successfully cleared {deleted_files} files'
        }
        
        if errors:
            response['warnings'] = errors
        
        return jsonify(response)
        
    except Exception as e:
        return jsonify({'error': f'Clear operation failed: {str(e)}'}), 500

@app.route('/cleanup')
def manual_cleanup():
    try:
        cleanup_old_files()
        return jsonify({'message': 'Manual cleanup completed successfully'})
    except Exception as e:
        return jsonify({'error': f'Cleanup failed: {str(e)}'}), 500

@app.route('/status')
def status():
    try:
        upload_count = len([f for f in os.listdir(app.config['UPLOAD_FOLDER']) if os.path.isfile(os.path.join(app.config['UPLOAD_FOLDER'], f))]) if os.path.exists(app.config['UPLOAD_FOLDER']) else 0
        output_count = len([f for f in os.listdir(app.config['OUTPUT_FOLDER']) if os.path.isfile(os.path.join(app.config['OUTPUT_FOLDER'], f))]) if os.path.exists(app.config['OUTPUT_FOLDER']) else 0
        
        return jsonify({
            'status': 'healthy',
            'upload_folder': app.config['UPLOAD_FOLDER'],
            'output_folder': app.config['OUTPUT_FOLDER'],
            'uploaded_files': upload_count,
            'output_files': output_count,
            'supported_formats': list(ALLOWED_EXTENSIONS),
            'max_file_size': app.config['MAX_CONTENT_LENGTH']
        })
    except Exception as e:
        return jsonify({'error': f'Status check failed: {str(e)}'}), 500

if __name__ == '__main__':
    print("Starting Image Converter Pro...")
    print(f"Upload folder: {app.config['UPLOAD_FOLDER']}")
    print(f"Output folder: {app.config['OUTPUT_FOLDER']}")
    print("Server starting on http://0.0.0.0:5000")
    
    try:
        app.run(debug=True, host='0.0.0.0', port=5000)
    except Exception as e:
        print(f"Failed to start server: {e}")
